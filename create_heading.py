import pandas as pd
from docx import Document


def crear_docx_por_fila(archivo_excel, columna_nombre, carpeta_destino):
    # Leer el archivo Excel
    df = pd.read_excel(archivo_excel)

    # Crear un documento .docx para cada fila
    for index, row in df.iterrows():
        # Crear un nuevo documento
        doc = Document()

        # Agregar contenido al documento con los datos de la fila
        doc.add_heading(f"Documento para {row[columna_nombre]}", level=1)
        doc.add_heading(f"Documento para {row[columna_nombre]}", level=2)

        for col in df.columns:
            doc.add_paragraph(f"{col}: {row[col]}")

        # Guardar el documento .docx con el nombre de la columna indicada
        nombre_archivo = f"{carpeta_destino}/{row[columna_nombre]}.docx"
        doc.save(nombre_archivo)
        print(f"Documento creado: {nombre_archivo}")


# Usar la función
archivo_excel = 'plantilla_alumnos.xlsx'  # Cambia esto a la ruta de tu archivo Excel
columna_nombre = 'Nombre'  # Columna que se usará para nombrar cada documento
carpeta_destino = '/home/jona/PycharmProjects/Word_pandas'  # Carpeta donde se guardarán los documentos

crear_docx_por_fila(archivo_excel, columna_nombre, carpeta_destino)
